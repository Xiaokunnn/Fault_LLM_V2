#!/usr/bin/env python3
"""Build the RP2 Chinese review draft in an IEEE-conference-like DOCX layout."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "papers/D2AI_ICDM_2026/RP2_D2AI2026_中文初稿_v3.md"
DEFAULT_OUTPUT = ROOT / "papers/D2AI_ICDM_2026/RP2_D2AI2026_中文初稿_v3.docx"

LATIN_FONT = "Times New Roman"
CN_FONT = "Noto Serif SC"
MATH_FONT = "Cambria Math"


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False,
                 latin: str = LATIN_FONT, east_asia: str = CN_FONT) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), latin)


def set_style_font(style, size: float, *, bold: bool = False, italic: bool = False) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), CN_FONT)
    fonts.set(qn("w:cs"), LATIN_FONT)


def set_columns(section, count: int, space_twips: int = 240) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        col = cols[0]
    else:
        col = OxmlElement("w:cols")
        sect_pr.append(col)
    col.set(qn("w:num"), str(count))
    col.set(qn("w:space"), str(space_twips))
    col.set(qn("w:equalWidth"), "1")


def configure_section(section, columns: int) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.667)
    section.right_margin = Inches(0.667)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.35)
    set_columns(section, columns)


def keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_inline(paragraph, text: str, size: float, *, default_bold: bool = False,
               default_italic: bool = False) -> None:
    text = normalize_math(text)
    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`|\*[^*]+?\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size, bold=default_bold, italic=default_italic)
        token = match.group(0)
        if token.startswith("**"):
            content, bold, italic, latin = token[2:-2], True, default_italic, LATIN_FONT
        elif token.startswith("`"):
            content, bold, italic, latin = token[1:-1], default_bold, False, "Consolas"
        else:
            content, bold, italic, latin = token[1:-1], default_bold, True, LATIN_FONT
        run = paragraph.add_run(content)
        set_run_font(run, size, bold=bold, italic=italic, latin=latin)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size, bold=default_bold, italic=default_italic)


def normalize_math(text: str) -> str:
    text = re.sub(r"\\tag\{([^{}]+)\}", r"(\1)", text)
    replacements = {
        r"\(": "", r"\)": "", r"\[": "", r"\]": "",
        r"\mathcal G": "G", r"\mathcal V": "V", r"\mathcal C": "C",
        r"\mathcal A": "A", r"\mathcal S": "S", r"\mathcal P": "P", r"\in": "∈",
        r"\notin": "∉", r"\leq": "≤", r"\geq": "≥",
        r"\le": "≤", r"\ge": "≥", r"\cap": "∩",
        r"\lambda_o": "λ_o", r"\Delta": "Δ", r"\mid": "|",
        r"\Phi": "Φ", r"\phi": "φ", r"\min": "min",
        r"\mathrm": "", r"\mathbb I": "I", r"\text": "", r"\quad": " ", r"\,": " ",
        r"\\": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", text).strip()


def add_body_paragraph(doc: Document, text: str, *, size: float = 9.3,
                       indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(size) if indent else Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(1.3)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(10.8)
    add_inline(paragraph, text, size)


def add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(5.5)
    fmt.space_after = Pt(2.0)
    fmt.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, 9.5, bold=False)
    run.font.small_caps = True


def add_subheading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(3.5)
    fmt.space_after = Pt(1.0)
    fmt.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, 9.3, italic=True)


def add_equation(doc: Document, lines: list[str]) -> None:
    raw = " ".join(line.strip() for line in lines if line.strip())
    if "begin{cases}" in raw:
        display = "U_q = F1_q（可回答且契约有效）；1（不可回答且合法拒答）；0（其他）。  (3)"
    else:
        display = normalize_math(raw)
        display = display.replace(", (1)", ".  (1)")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(1.5)
    fmt.space_after = Pt(1.5)
    fmt.keep_together = True
    run = paragraph.add_run(display)
    set_run_font(run, 8.2, italic=True, latin=MATH_FONT)


def set_cell_margins(cell, top: int = 30, start: int = 40, bottom: int = 30,
                     end: int = 40) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_table_geometry(table, widths_inches: list[float]) -> None:
    total = int(round(sum(widths_inches) * 1440))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = int(round(widths_inches[idx] * 1440))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    if cols == 6:
        widths = [0.66, 0.36, 0.36, 0.84, 0.72, 0.54]
    elif cols == 5:
        widths = [0.91, 0.62, 0.76, 0.45, 0.74]
    else:
        widths = [3.48 / cols] * cols
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(7.6)
            add_inline(paragraph, value.replace("**", ""), 6.4,
                       default_bold=(r_idx == 0))
            if r_idx == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "E7E6E6")
                cell._tc.get_or_add_tcPr().append(shd)
    configure_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)


def add_figure(doc: Document, image_path: Path, caption: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    # Normal uses exact 10.8 pt leading. Clear it here so Word does not clip the image
    # to a single text line when the figure is placed in a narrow IEEE column.
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run()
    picture = run.add_picture(str(image_path), width=Inches(3.34))
    alt_text = (caption or image_path.stem).replace("**", "").strip()
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(2.5)
        cap.paragraph_format.keep_with_next = False
        add_inline(cap, caption.replace("**", ""), 7.3)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[str] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw.append(lines[idx].strip())
        idx += 1
    rows = []
    for row in raw:
        cells = [cell.strip() for cell in row.strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def build(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.core_properties.title = lines[0].removeprefix("# ").strip()
    doc.core_properties.subject = "D2AI 2026 Chinese review draft"
    doc.core_properties.author = "Author information pending"
    doc.core_properties.comments = "RP2 v6 Chinese review draft. All evaluation labels are Silver; no domain-expert review."

    normal = doc.styles["Normal"]
    set_style_font(normal, 9.3)
    normal.paragraph_format.space_after = Pt(1.3)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(10.8)
    configure_section(doc.sections[0], 1)

    title = lines[0].removeprefix("# ").strip()
    english_title = lines[2].replace("**", "").strip()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(3)
    run = title_p.add_run(title)
    set_run_font(run, 18.2)
    english_p = doc.add_paragraph()
    english_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    english_p.paragraph_format.space_before = Pt(0)
    english_p.paragraph_format.space_after = Pt(5)
    run = english_p.add_run(english_title)
    set_run_font(run, 9.0, italic=True)

    author_lines = [
        line.replace("**作者：**", "").strip().rstrip("  ")
        for line in lines if line.startswith("**作者：**")
    ]
    affiliation_lines = [
        line.replace("**单位：**", "").strip().rstrip("  ")
        for line in lines if line.startswith("**单位：**")
    ]
    email_lines = [
        line.replace("**邮箱：**", "").strip()
        for line in lines if line.startswith("**邮箱：**")
    ]
    for content, size, italic in (
        ((author_lines or ["[作者姓名]，[导师/合作者姓名]"])[0], 10.0, False),
        ((affiliation_lines or ["[学院、学校，城市，国家]"])[0], 8.5, True),
        ((email_lines or ["[通信作者邮箱]"])[0], 8.3, False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0.5)
        run = p.add_run(content)
        set_run_font(run, size, italic=italic)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, 2)

    start = next(i for i, line in enumerate(lines) if line.strip() == "## 摘要")
    idx = start
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue
        if line.startswith(">") or line.startswith("**作者：**") or line.startswith("**单位：**") or line.startswith("**邮箱：**"):
            idx += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading in {"摘要", "Abstract"}:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(1)
                label = "摘要—" if heading == "摘要" else "Abstract—"
                r = p.add_run(label)
                set_run_font(r, 8.3, bold=True, italic=True)
                idx += 1
                while idx < len(lines) and not lines[idx].strip():
                    idx += 1
                if idx < len(lines):
                    add_inline(p, lines[idx].strip(), 8.3)
                    idx += 1
                continue
            add_section_heading(doc, heading)
            idx += 1
            continue
        if line.startswith("### "):
            add_subheading(doc, line[4:].strip())
            idx += 1
            continue
        if line.startswith("**关键词：**") or line.startswith("**Index Terms"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, line, 8.3)
            idx += 1
            continue
        if line == r"\[":
            eq_lines = []
            idx += 1
            while idx < len(lines) and lines[idx].strip() != r"\]":
                eq_lines.append(lines[idx])
                idx += 1
            add_equation(doc, eq_lines)
            idx += 1
            continue
        if line.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            image_path = (source.parent / image_match.group(2)).resolve()
            caption = None
            lookahead = idx + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            if lookahead < len(lines) and lines[lookahead].strip().startswith("**图"):
                caption = lines[lookahead].strip()
                idx = lookahead + 1
            else:
                idx += 1
            add_figure(doc, image_path, caption)
            continue
        if line.startswith("[1]") or re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.first_line_indent = Pt(-10)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1.0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(8.3)
            add_inline(p, line, 7.2)
            idx += 1
            continue

        paragraph_lines = [line]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].strip()
            if not nxt:
                break
            if (nxt.startswith("## ") or nxt.startswith("### ") or nxt.startswith("|")
                    or nxt.startswith("![") or nxt == r"\[" or nxt.startswith("[1]")
                    or re.match(r"^\[\d+\]", nxt)):
                break
            paragraph_lines.append(nxt)
            idx += 1
        text = " ".join(paragraph_lines)
        is_declaration = any(marker in text for marker in (
            "Conceptualization", "全体作者声明", "生成式AI工具用于", "冻结配置、代码",
            "本研究使用公开技术文档",
        ))
        add_body_paragraph(doc, text, size=7.8 if is_declaration else 9.3,
                           indent=not is_declaration)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Created {output}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
